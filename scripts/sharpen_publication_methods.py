from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph


METHODS: tuple[tuple[str, str], ...] = (
    (
        "3.1 Study design and evaluation data",
        "The workflow was evaluated as a local, fail-closed de-identification pipeline. "
        "The end-to-end experiment used all 23,921 DICOM instances in the MIDI-B "
        "Synthetic Validation collection. A localization analysis used the 35 manually "
        "annotated burned-in-PHI regions available in that collection. The detector had "
        "been refined using this validation split; consequently, localization and official "
        "DICOM results are reported as development and pre-production validation, not as "
        "performance on an independent holdout. Source DICOM data, mapping files, model "
        "artifacts, and validator output remained on the controlled processing host.",
    ),
    (
        "3.2 Pixel-PHI localization and masking",
        "Every decoded image or video frame was passed through the configured PHI masking "
        "path. In video processing, RapidOCR supplied local text observations for metadata "
        "extraction. A checksum-pinned YOLOv8n ONNX detector supplied additive PHI-region "
        "proposals for video frames and DICOM images. "
        "Production inference used 960-pixel letterboxed input, confidence threshold 0.05, "
        "non-maximum-suppression threshold 0.45, and PHI class 0. Model absence, checksum "
        "mismatch, unreadable output, or an invalid image caused processing to fail rather "
        "than yield a releasable artifact. Detector boxes were transformed back to native "
        "image coordinates, clipped to image bounds, and filled with opaque black pixels. "
        "Device-specific static regions could be added, but neither detector nor OCR output "
        "could remove a deterministic mask. The refined localization analysis used "
        "confidence threshold 0.10 and is therefore reported separately from the 0.05 "
        "end-to-end export configuration.",
    ),
    (
        "3.3 Metadata transformation and pseudonymous linkage",
        "DICOM attributes were normalized once at ingestion and then processed by an "
        "explicit action profile for removal, replacement, retention, or consistency "
        "checking. Date candidates were extracted from one shared candidate set and "
        "assigned by label and timestamp context; dates were shifted by a stable offset per "
        "patient. Patient identifiers were replaced by deterministic study pseudonyms, and "
        "Study, Series, and SOP Instance UIDs were remapped while preserving internal "
        "references and file-meta consistency. Patient-ID and UID mappings were stored "
        "outside the exported DICOM tree for use by the official validator. These values "
        "support controlled longitudinal linkage and are treated as pseudonyms, not as "
        "evidence that the records are anonymous.",
    ),
    (
        "3.4 DICOM export and official validation",
        "Each transformed instance was written to a patient/study/series hierarchy and "
        "reloaded before inclusion in the final export. The completed tree was evaluated "
        "with the challenge organizers' MIDI validation code against "
        "MIDI-B-Answer-Key-Validation.db. The registered run "
        "MIDI_B_Synthetic_Validation_Preproduction_20260715 used eight worker processes, "
        "the exported patient-ID and UID mapping files, and database-backed result output. "
        "We recorded the number of indexed files, missing-file batches, database integrity, "
        "and every action-level pass or failure. No file was omitted from the reported "
        "end-to-end result.",
    ),
    (
        "3.5 Outcomes and statistical analysis",
        "For box localization, a prediction was a true positive when it matched an annotated "
        "PHI box at intersection over union (IoU) >= 0.50. We report true positives, "
        "precision TP/(TP+FP), recall TP/(TP+FN), F1, mean best IoU, and mean annotated-box "
        "coverage. Results were also stratified by modality (CR, DX, and MG). For DICOM "
        "validation, action pass rate was pass/(pass+fail). The raw pass rate pooled all "
        "actions. The official challenge score was the weighted sum of category pass rates: "
        "0.70 x HIPAA + 0.20 x DICOM Standard + 0.10 x TCIA Best Practice. Counts are "
        "reported without sampling confidence intervals because the validator exhaustively "
        "tested the exported collection; this does not remove uncertainty about "
        "generalization to other institutions, devices, or populations.",
    ),
    (
        "3.6 Release controls",
        "Automated processing did not authorize release. Approval remained disabled until a "
        "playable or reloadable anonymized artifact existed, model integrity checks had "
        "passed, and a qualified reviewer had inspected the complete output and recorded a "
        "decision. Rejected artifacts returned to correction and re-review. For structured "
        "secondary-use tables, quasi-identifiers, the minimum class size k, sensitive-"
        "attribute constraints, and the maximum permitted utility discrepancy were declared "
        "before release. A candidate view was accepted only if every full quasi-identifier "
        "class met k, every configured sensitive-attribute constraint passed, and measured "
        "utility discrepancy remained within its limit. Failure of any privacy, utility, "
        "provenance, or review condition yielded no release. This tabular control is distinct "
        "from the official DICOM score and does not convert pseudonymous records into "
        "anonymous data.",
    ),
)


def _find_exact(document: DocumentObject, text: str) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == text]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def _append_before(
    document: DocumentObject,
    anchor: Paragraph,
    text: str,
    style: str | None = None,
) -> Paragraph:
    paragraph = document.add_paragraph(text, style=style)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def _replace_methods(document: DocumentObject) -> None:
    start = _find_exact(document, "3. Methods")
    end = _find_exact(document, "5. Results")
    cursor = start._p.getnext()
    while cursor is not None and cursor is not end._p:
        following = cursor.getnext()
        cursor.getparent().remove(cursor)
        cursor = following

    _append_before(
        document,
        end,
        "We implemented a local pipeline that transforms visible PHI, DICOM metadata, and "
        "linkage identifiers before secondary use. The evaluated path comprised pixel-PHI "
        "localization, deterministic metadata transformation, pseudonymous identifier "
        "mapping, DICOM export, official validation, and mandatory review. Processing was "
        "fail closed: an incomplete transformation or failed integrity check produced no "
        "release candidate.",
    )
    for heading, body in METHODS:
        _append_before(document, end, heading, style="Heading 2")
        _append_before(document, end, body)


def _renumber_after_methods(document: DocumentObject) -> None:
    exact_replacements = {
        "5. Results": "4. Results",
        "5.1 Video-Level Anonymization Throughput and System Latency": (
            "4.1 Video-Level Anonymization Throughput and System Latency"
        ),
        "5.2 OCR and De-Identification Quality Evaluation": (
            "4.2 OCR and De-Identification Quality Evaluation"
        ),
        "Official end-to-end DICOM validation results": (
            "4.3 Official end-to-end DICOM validation results"
        ),
        "5.4 Verification of Privacy-Preserving Data Publishing Frameworks": (
            "4.4 Verification of Privacy-Preserving Data Publishing Frameworks"
        ),
        "6. Limitations": "5. Limitations",
        "7. Conclusion": "6. Conclusion",
    }
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value in exact_replacements:
            paragraph.text = exact_replacements[value]
        elif value.startswith("Figure 3. Official MIDI-B category"):
            paragraph.text = value.replace("Figure 3.", "Figure 1.", 1)
        elif value.startswith("Figure 4. Official action-level"):
            paragraph.text = value.replace("Figure 4.", "Figure 2.", 1)
        elif "Section 5.2 reports" in value:
            paragraph.text = value.replace("Section 5.2 reports", "Section 4.2 reports")


def sharpen_methods(source: Path, output: Path) -> None:
    if source.resolve() == output.resolve():
        raise ValueError("Source and output paths must differ")
    document = Document(str(source))
    _replace_methods(document)
    _renumber_after_methods(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Consolidate the manuscript Methods section")
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    return parser


def main() -> int:
    arguments = _parser().parse_args()
    sharpen_methods(arguments.source, arguments.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
