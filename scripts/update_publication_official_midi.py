from __future__ import annotations

import argparse
import subprocess
import tempfile
from pathlib import Path
from typing import Protocol, cast

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.font import Font
from docx.text.paragraph import Paragraph


class _StyleWithFont(Protocol):
    font: Font


WEIGHTED_ROWS = (
    ("HIPAA", 520_616, 47_243, 91.6805, 70, 64.1764),
    ("DICOM Standard", 1_975_772, 2_197, 99.8889, 20, 19.9778),
    ("Best Practice", 2_168_258, 1_459_118, 59.7748, 10, 5.9775),
)

ACTION_ROWS = (
    ("Date shifted", 139_774, 0, 100.0000),
    ("Patient ID consistent", 23_921, 0, 100.0000),
    ("Pixels hidden", 35, 0, 100.0000),
    ("Pixels retained", 23_884, 2, 99.9916),
    ("Tag retained", 1_098_954, 2_137, 99.8059),
    ("Text non-null", 618_481, 58, 99.9906),
    ("Text removed", 218_482, 125_101, 63.5893),
    ("Text retained", 2_072_283, 1_381_256, 60.0046),
    ("UID changed", 234_416, 2, 99.9991),
    ("UID consistent", 234_416, 2, 99.9991),
)

ABSTRACT_RESULT = (
    "For the release candidate evaluated here, the production video path combines "
    "the local RapidOCR cascade, deterministic metadata extraction, an additive "
    "ONNX PHI-region detector, fail-closed model integrity verification, and mandatory "
    "human approval before secondary-use release. In the refined detector-development "
    "evaluation on 35 human-annotated burned-in-PHI instances from MIDI-B Synthetic "
    "Validation, the detector matched all 35 annotations from 44 predictions "
    "(precision 0.795, recall 1.000, F1 0.886, mean best IoU 0.897; matching IoU 0.50). "
    "The complete 23,921-file collection was subsequently exported as a pseudonymized "
    "DICOM tree and processed by the challenge organizers’ validation code. All 35 "
    "pixels-hidden checks, all 139,774 date-shift checks, and all 23,921 patient-ID "
    "consistency checks passed. Across 6,173,204 checks, the raw pass rate was 75.56% "
    "and the organizer-weighted score was 90.13% (HIPAA 70%, DICOM Standard 20%, "
    "best practice 10%). The principal residual deficits concerned text removal and "
    "retention policy. Because the detector was refined using this validation split, "
    "these results are development and pre-production validation evidence, not an "
    "untouched holdout or autonomous-release claim."
)

DETECTOR_METHOD = (
    "The refined pinned ONNX detector was executed directly on the human-annotated "
    "DICOM instances in the MIDI-B Synthetic Validation collection. The detector "
    "development evaluation used confidence 0.10, 960-pixel letterboxed input, and "
    "one-to-one box matching at IoU 0.50. All 35 annotated instances were evaluated "
    "without missing or failed cases. The later end-to-end DICOM export used the "
    "release confidence threshold of 0.05."
)

DETECTOR_INTERPRETATION = (
    "Mean ground-truth box coverage was 0.945. At this operating point no annotated "
    "burned-in-PHI box was missed among the 35 scored instances. The result supports "
    "the detector’s intended role as a sensitive, additive mask-proposal mechanism. "
    "It does not estimate specificity on unannotated negative regions, and some scored "
    "instances contributed to model refinement; it is therefore an operational "
    "development check rather than an external generalization estimate."
)

OFFICIAL_INTRO = (
    "The completed MIDI-B Synthetic Validation collection contained 23,921 DICOM "
    "files, matching all 23,921 answer-key records. A deterministic HMAC-based mapping "
    "boundary produced 216 patient-ID mappings and 27,578 UID mappings. The exporter "
    "recursively removed private attributes, transformed patient identifiers and UIDs, "
    "shifted dates consistently per patient, synchronized file-meta and dataset SOP "
    "Instance UIDs, preserved source PixelData when no mask was required, and applied "
    "photometrically correct black boxes to detector regions. It wrote 23,921 reloadable "
    "DICOM files; 37 files contained 56 masked detector regions."
)

LIMITATIONS = (
    "The end-to-end DICOM assessment used the MIDI-B Synthetic Validation split rather "
    "than an untouched test or external clinical cohort. The official validator "
    "completed 6,173,204 checks with no missing files, but the raw pass rate (75.56%) "
    "shows that conformance is not uniform across action classes. Pixel hiding, date "
    "shifting, and patient-ID consistency each passed 100%; text-removed and "
    "text-retained actions passed 63.59% and 60.00%, respectively. Many retention "
    "failures arise from the conservative removal of private attributes, while removal "
    "failures identify descriptive metadata that the current profile retains. Two "
    "pixels-retained checks also failed, indicating a small false-positive masking "
    "burden. These results justify targeted metadata-profile refinement and continued "
    "mandatory review; they do not support autonomous release or universal DICOM safety."
)

CONCLUSION = (
    "Absolute anonymity is not inferred from empirical similarity or satisfaction of "
    "configured privacy criteria. In the post-fix integration fixture, the production "
    "video OCR paths recovered all expected PHI metadata fields in four correlated "
    "frames. The refined detector matched all 35 annotated MIDI-B burned-in-PHI boxes "
    "from 44 predictions (precision 0.795, recall 1.000, F1 0.886), and the end-to-end "
    "DICOM export passed all official pixel-hidden, date-shift, and patient-ID "
    "consistency checks. The organizer-weighted score was 90.13% across 6,173,204 "
    "checks. Together with checksum-pinned fail-closed inference, local OCR fallback, "
    "controlled pseudonymous mappings, and mandatory review, these controls support a "
    "release-ready human-supervised workflow. Remaining metadata-policy failures, the "
    "two pixel-retention failures, validation-split refinement, and the absence of an "
    "external holdout preclude claims of autonomous safety or general clinical "
    "performance."
)


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("figures", type=Path)
    return parser


def _repair_docx(source: Path, destination: Path) -> None:
    destination.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(destination),
            str(source),
        ],
        check=True,
        capture_output=True,
        text=True,
    )


def _replace_start(document: DocumentObject, prefix: str, value: str) -> Paragraph:
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"expected one paragraph beginning {prefix!r}, got {len(matches)}")
    paragraph = matches[0]
    paragraph.text = value
    return paragraph


def _insert_paragraph_after(
    document: DocumentObject,
    anchor: Paragraph,
    text: str = "",
    *,
    style: str | None = None,
) -> Paragraph:
    paragraph = document.add_paragraph(text, style=style)
    anchor._p.addnext(paragraph._p)
    return paragraph


def _insert_table_after(
    document: DocumentObject,
    anchor: Paragraph,
    headers: tuple[str, ...],
    rows: tuple[tuple[str, ...], ...],
) -> tuple[Table, Paragraph]:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            cell.text = value
    anchor._p.addnext(table._tbl)
    trailing = document.add_paragraph()
    table._tbl.addnext(trailing._p)
    return table, trailing


def _add_picture_after(
    document: DocumentObject,
    anchor: Paragraph,
    image: Path,
    caption: str,
) -> Paragraph:
    image_paragraph = _insert_paragraph_after(document, anchor)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(image), width=Inches(6.25))
    caption_paragraph = _insert_paragraph_after(
        document,
        image_paragraph,
        caption,
        style="Caption",
    )
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return caption_paragraph


def _format_document(document: DocumentObject) -> None:
    normal = cast(_StyleWithFont, document.styles["Normal"])
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)


def _weighted_figure(path: Path) -> None:
    names = [row[0] for row in WEIGHTED_ROWS]
    rates = [row[3] for row in WEIGHTED_ROWS]
    contributions = [row[5] for row in WEIGHTED_ROWS]
    colors = ["#0072B2", "#009E73", "#E69F00"]
    fig, axes = plt.subplots(1, 2, figsize=(10.2, 4.1), constrained_layout=True)
    y = np.arange(len(names))
    axes[0].barh(y, rates, color=colors, height=0.58)
    axes[0].set_yticks(y, names)
    axes[0].invert_yaxis()
    axes[0].set_xlim(0, 105)
    axes[0].set_xlabel("Pass rate (%)")
    axes[0].set_title("Official category pass rates")
    for index, value in enumerate(rates):
        axes[0].text(value + 1, index, f"{value:.2f}%", va="center", fontsize=9)

    axes[1].barh(y, contributions, color=colors, height=0.58)
    axes[1].set_yticks(y, names)
    axes[1].invert_yaxis()
    axes[1].set_xlim(0, 72)
    axes[1].set_xlabel("Weighted contribution (percentage points)")
    axes[1].set_title("Contribution to weighted score")
    for index, value in enumerate(contributions):
        axes[1].text(value + 0.7, index, f"{value:.2f}", va="center", fontsize=9)
    axes[1].text(
        0.98,
        0.05,
        "Weighted score\n90.13%",
        transform=axes[1].transAxes,
        ha="right",
        va="bottom",
        fontsize=13,
        fontweight="bold",
        color="#333333",
    )
    for axis in axes:
        axis.spines[["top", "right"]].set_visible(False)
        axis.grid(axis="x", alpha=0.2)
    fig.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def _action_figure(path: Path) -> None:
    names = [row[0] for row in ACTION_ROWS]
    rates = np.asarray([row[3] for row in ACTION_ROWS])
    failures = np.asarray([row[2] for row in ACTION_ROWS])
    colors = ["#009E73" if value >= 99.9 else "#D55E00" for value in rates]
    y = np.arange(len(names))
    fig, axes = plt.subplots(
        1,
        2,
        figsize=(10.4, 6.0),
        gridspec_kw={"width_ratios": (1.45, 1)},
        constrained_layout=True,
    )
    axes[0].barh(y, rates, color=colors, height=0.62)
    axes[0].set_yticks(y, names)
    axes[0].invert_yaxis()
    axes[0].set_xlim(55, 101.5)
    axes[0].set_xlabel("Pass rate (%) — axis begins at 55%")
    axes[0].set_title("Action-level pass rates")
    for index, value in enumerate(rates):
        axes[0].text(min(value + 0.45, 100.2), index, f"{value:.2f}%", va="center", fontsize=8)

    display_failures = np.maximum(failures, 0.5)
    axes[1].barh(y, display_failures, color="#CC79A7", height=0.62)
    axes[1].set_yticks(y, [""] * len(names))
    axes[1].invert_yaxis()
    axes[1].set_xscale("log")
    axes[1].set_xlabel("Failed checks (log scale)")
    axes[1].set_title("Failure counts")
    for index, value in enumerate(failures):
        axes[1].text(max(value, 0.5) * 1.15, index, f"{value:,}", va="center", fontsize=8)
    for axis in axes:
        axis.spines[["top", "right"]].set_visible(False)
        axis.grid(axis="x", alpha=0.2)
    fig.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def _update_detector_table_paragraphs(document: DocumentObject) -> None:
    replacements = {
        "16": "15",
        "0.875": "0.933",
        "0.933": "0.966",
        "30": "24",
        "0.600": "0.750",
        "0.750": "0.857",
        "6": "5",
        "0.500": "0.600",
        "0.667": "0.750",
        "52": "44",
        "0.673": "0.795",
        "0.805": "0.886",
    }
    modality_anchor = next(
        index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "Modality"
    )
    for paragraph in document.paragraphs[modality_anchor : modality_anchor + 40]:
        value = paragraph.text.strip()
        if value in replacements:
            paragraph.text = replacements[value]


def update_document(source: Path, output: Path, figures: Path) -> None:
    figures.mkdir(parents=True, exist_ok=True)
    weighted_figure = figures / "midi_b_official_weighted_categories.png"
    action_figure = figures / "midi_b_official_action_results.png"
    _weighted_figure(weighted_figure)
    _action_figure(action_figure)

    with tempfile.TemporaryDirectory(prefix="lx-paper-repair-") as temp_directory:
        repair_root = Path(temp_directory)
        _repair_docx(source, repair_root)
        repaired = repair_root / source.name
        document = Document(str(repaired))

        _format_document(document)
        _replace_start(document, "For the release candidate evaluated here", ABSTRACT_RESULT)
        _replace_start(document, "The pinned ONNX detector was executed", DETECTOR_METHOD)
        _replace_start(document, "The mean ground-truth box coverage", DETECTOR_INTERPRETATION)
        official_anchor = _replace_start(
            document,
            "The installed collection contained",
            OFFICIAL_INTRO,
        )
        _replace_start(document, "The measured DICOM result is limited", LIMITATIONS)
        _replace_start(document, "Absolute anonymity is not inferred", CONCLUSION)
        _update_detector_table_paragraphs(document)

        anchor = _insert_paragraph_after(
            document,
            official_anchor,
            "Official end-to-end DICOM validation results",
            style="Heading 3",
        )
        anchor = _insert_paragraph_after(
            document,
            anchor,
            "The organizer validator indexed all exported files, reported zero missing-file "
            "batches, completed 6,173,204 checks, and produced a database that passed SQLite "
            "integrity verification. The raw pass rate was 75.56%. Applying the organizer’s "
            "published category weights yielded a challenge score of 90.13%.",
        )
        weighted_table_rows = tuple(
            (
                name,
                f"{passed:,}",
                f"{failed:,}",
                f"{rate:.4f}%",
                f"{weight}%",
                f"{contribution:.4f}%",
            )
            for name, passed, failed, rate, weight, contribution in WEIGHTED_ROWS
        ) + (("Weighted total", "4,664,646", "1,508,558", "75.5628% raw", "100%", "90.1316%"),)
        _, anchor = _insert_table_after(
            document,
            anchor,
            ("Category", "Pass", "Fail", "Pass rate", "Weight", "Contribution"),
            weighted_table_rows,
        )
        anchor = _add_picture_after(
            document,
            anchor,
            weighted_figure,
            "Figure 3. Official MIDI-B category pass rates and weighted contributions. "
            "The 90.13% score uses HIPAA, DICOM Standard, and best-practice weights of "
            "70%, 20%, and 10%, respectively.",
        )
        action_table_rows = tuple(
            (name, f"{passed:,}", f"{failed:,}", f"{passed + failed:,}", f"{rate:.4f}%")
            for name, passed, failed, rate in ACTION_ROWS
        )
        _, anchor = _insert_table_after(
            document,
            anchor,
            ("Official action", "Pass", "Fail", "Total", "Pass rate"),
            action_table_rows,
        )
        anchor = _add_picture_after(
            document,
            anchor,
            action_figure,
            "Figure 4. Official action-level pass rates and failure counts. The pass-rate "
            "axis is truncated at 55% for readability; failure counts use a logarithmic scale.",
        )
        _insert_paragraph_after(
            document,
            anchor,
            "Privacy-critical checks were strongest: pixels hidden (35/35), dates shifted "
            "(139,774/139,774), and patient-ID consistency (23,921/23,921) all passed. "
            "Pixels retained passed 23,884/23,886 checks. The dominant deficits were text "
            "removal (63.59%) and text retention (60.00%), showing that the metadata profile "
            "requires further policy alignment even though identifier mapping and pixel "
            "hiding were highly reliable.",
        )

        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output))


def main() -> int:
    args = _parser().parse_args()
    update_document(args.source.resolve(), args.output.resolve(), args.figures.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
