from __future__ import annotations

import argparse
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph

from scripts.convert_publication_formulas_to_omml import Math, Part, set_word_display_math
from scripts.convert_publication_formulas_to_omml import set_word_math


ABSTRACT = (
    "Endoscopic media can retain protected health information (PHI) in both metadata and "
    "image pixels. We developed a device-aware, locally operated workflow that combines "
    "temporal metadata accumulation, deterministic pseudonymous mapping, additive PHI-region "
    "detection, opaque masking, fail-closed artifact generation, and mandatory human review. "
    "The primary contribution is the reproducible processing and release-control method; a "
    "synthetic-row mechanism is retained only as an experimental released-view control and "
    "is not presented as person-level anonymity."
)

ABSTRACT_RESULTS = (
    "A four-frame, single-video test verified execution of the corrected endoscopy OCR path "
    "but was not statistically independent and does not estimate clinical reliability. In a "
    "separate detector-development analysis of 35 annotated MIDI-B regions, all annotations "
    "were matched from 44 predictions (precision 0.795, recall 1.000, F1 0.886, mean best IoU "
    "0.897). The shared DICOM export and masking boundary was then applied to 23,921 CR, DX, "
    "and MG instances and evaluated with the challenge organizers' validator. Across "
    "6,173,204 correlated action checks, the raw pass rate was 75.56% and the fixed-weight "
    "challenge score was 90.13%. Pixel hiding, date shifting, and patient-ID consistency "
    "passed all applicable checks, whereas text removal passed 63.59% and text retention "
    "60.00%."
)

ABSTRACT_LIMITS = (
    "These findings establish technical integration and pre-production DICOM conformance, not "
    "independent endoscopy-video safety. The detector was refined on the validation split, the "
    "four video frames came from one patient, and failed text-removal checks require "
    "attribute-level review before release. A frozen patient- and video-level test set across "
    "devices, sites, PHI types, and image conditions remains necessary."
)


def _exact(document: DocumentObject, text: str) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == text]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def _starts(document: DocumentObject, text: str) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith(text)]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph starting {text!r}, found {len(matches)}")
    return matches[0]


def _index(paragraphs: list[Paragraph], target: Paragraph) -> int:
    for index, paragraph in enumerate(paragraphs):
        if paragraph._p is target._p:
            return index
    raise ValueError(f"Paragraph not found: {target.text!r}")


def _remove_paragraphs(paragraphs: Iterable[Paragraph]) -> None:
    for paragraph in paragraphs:
        parent = paragraph._p.getparent()
        if parent is not None:
            parent.remove(paragraph._p)


def _new_paragraph(
    document: DocumentObject,
    text: str,
    *,
    style: str | None = None,
) -> Paragraph:
    return document.add_paragraph(text, style=style)


def _insert_after(anchor: Paragraph, elements: Sequence[Paragraph | Table]) -> None:
    cursor = anchor._p
    for element in elements:
        xml = element._p if isinstance(element, Paragraph) else element._tbl
        cursor.addnext(xml)
        cursor = xml


def _caption(document: DocumentObject, text: str) -> Paragraph:
    paragraph = document.add_paragraph(text, style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.right_indent = Inches(0)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def _shade_cell(cell: _Cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _repeat_header(row: _Row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _prevent_row_split(row: _Row) -> None:
    properties = row._tr.get_or_add_trPr()
    no_split = OxmlElement("w:cantSplit")
    properties.append(no_split)


def _format_table(
    table: Table,
    widths: Sequence[float],
    *,
    numeric_from: int = 1,
) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        _prevent_row_split(row)
        for column_index, cell in enumerate(row.cells):
            cell.width = Inches(widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _shade_cell(cell, "1F4E78" if row_index == 0 else ("D9EAF7" if row_index % 2 == 0 else "FFFFFF"))
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                if row_index == 0 or column_index >= numeric_from:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5)
                    if row_index == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def _table(
    document: DocumentObject,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[float],
    *,
    numeric_from: int = 1,
) -> Table:
    table = document.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            cell.text = value
    _format_table(table, widths, numeric_from=numeric_from)
    return table


def _rewrite_abstract_and_focus(document: DocumentObject) -> None:
    abstract_heading = _exact(document, "Abstract")
    introduction = _exact(document, "1. Introduction")
    paragraphs = document.paragraphs
    start = _index(paragraphs, abstract_heading)
    stop = _index(paragraphs, introduction)
    old = paragraphs[start + 1 : stop]
    replacements = (ABSTRACT, ABSTRACT_RESULTS, ABSTRACT_LIMITS)
    for paragraph, value in zip(old, replacements, strict=False):
        paragraph.text = value
        paragraph.style = "Body Text"
    _remove_paragraphs(old[len(replacements) :])

    focus = _starts(document, "The clinical problem is therefore not simply")
    focus.text = (
        "The primary research question is whether a device-aware, locally operated workflow "
        "can detect and mask identifying information in heterogeneous clinical media while "
        "maintaining auditable pseudonymous linkage and release control. This paper therefore "
        "treats the de-identification workflow as its principal contribution. The tabular "
        "synthetic-row mechanism is secondary and experimental; it evaluates a declared "
        "released-view predicate and does not establish person-level anonymity."
    )


def _add_threat_model(document: DocumentObject) -> None:
    renames = {
        "3.2 Pixel-PHI localization and masking": "3.3 Pixel-PHI localization and masking",
        "3.3 Metadata transformation and pseudonymous linkage": (
            "3.4 Metadata transformation and pseudonymous linkage"
        ),
        "3.4 DICOM export and official validation": "3.5 DICOM export and official validation",
        "3.5 Outcomes and statistical analysis": "3.6 Outcomes and statistical analysis",
        "3.6 Release controls": "3.7 Release controls",
    }
    for paragraph in document.paragraphs:
        if paragraph.text.strip() in renames:
            paragraph.text = renames[paragraph.text.strip()]

    anchor = _starts(document, "The workflow was evaluated as a local, fail-closed")
    heading = _new_paragraph(document, "3.2 Threat model and intended guarantee", style="Heading 2")
    body = _new_paragraph(
        document,
        "The release scenario is an authorized research consortium operating through a "
        "controlled analysis environment. The model addresses linkage, membership, and "
        "synthetic-provenance inference by a recipient with plausible clinical background "
        "knowledge. It does not cover a compromised processing host, malicious custodian, "
        "stolen key material, or unrestricted auxiliary information. The intended guarantee "
        "is risk reduction and predicate-compliant release, not anonymity.",
    )
    caption = _caption(document, "Table 1. Threat model and scope of the claimed guarantee.")
    table = _table(
        document,
        ("Element", "Study assumption"),
        (
            ("Recipient", "Authorized consortium analyst in a controlled research environment"),
            ("Auxiliary knowledge", "Center, date range, procedure, diagnosis, and cohort membership"),
            ("Visible data", "Approved processed media, DICOM attributes, quasi-identifiers, and pseudonyms"),
            ("Protected data", "Direct mappings, source media, keys, and identifiable audit records"),
            ("Trusted parties", "Local custodian, processing host, key service, checker, and human reviewer"),
            ("Attacker capabilities", "Linkage, membership inference, and synthetic-provenance inference"),
            ("Excluded threats", "Compromised endpoint, malicious custodian, stolen keys, arbitrary side information"),
            ("Intended guarantee", "De-identification risk reduction and predicate-compliant release"),
        ),
        (1.35, 5.15),
        numeric_from=2,
    )
    _insert_after(anchor, (heading, body, caption, table))


def _clarify_methods_and_endpoints(document: DocumentObject) -> None:
    dicom = _starts(document, "Each transformed instance was written")
    dicom.text = (
        "The MIDI-B Synthetic Validation collection comprised CR, DX, and MG DICOM images, "
        "not endoscopic video or converted video frames. We used it to test the workflow's "
        "shared DICOM boundary: identifier mapping, date shifting, UID remapping, pixel-PHI "
        "detection, masking, export, and readback. Each transformed instance was written to a "
        "patient/study/series hierarchy and reloaded. The completed tree was evaluated against "
        "MIDI-B-Answer-Key-Validation.db with the challenge organizers' validation code. The "
        "registered run MIDI_B_Synthetic_Validation_Preproduction_20260715 used eight workers "
        "and the exported patient-ID and UID mapping files. The validator's published category "
        "weights were fixed by the organizers. This experiment assesses DICOM transformation "
        "and conformance; it does not test temporal accumulation or endoscopy-video "
        "generalization."
    )

    outcomes = _starts(document, "For box localization, a prediction was a true positive")
    set_word_math(
        outcomes,
        (
            "Extraction and redaction were evaluated separately. PHI-field recall measures "
            "correct structured extraction; it does not establish that pixels are safe. For "
            "region localization, a prediction was a true positive at ",
            Math(r"\operatorname{IoU}\ge0.50"),
            ", with precision ",
            Math(r"P=\frac{TP}{TP+FP}"),
            ", recall ",
            Math(r"R=\frac{TP}{TP+FN}"),
            ", and ",
            Math(r"F_1=\frac{2PR}{P+R}"),
            ". Safety-oriented endpoints are residual readable-PHI rate after masking, "
            "frame- and video-level complete-redaction rate, false-positive masked-area "
            "fraction, clinically relevant viewport loss, and manual-review rejection rate. "
            "The available data estimate localization metrics but do not estimate all of "
            "these safety endpoints. DICOM action pass rate was ",
            Math(r"r=\frac{\mathrm{pass}}{\mathrm{pass}+\mathrm{fail}}"),
            ". The weighted challenge score was ",
            Math(
                r"S=0.70r_{\mathrm{HIPAA}}+0.20r_{\mathrm{DICOM}}+"
                r"0.10r_{\mathrm{BestPractice}}"
            ),
            ". Checks within a file are correlated and are therefore reported as exhaustive "
            "conformance counts, not independent statistical samples.",
        ),
    )
    spatial = _new_paragraph(
        document,
        "For reference boxes Gi and predicted boxes Pj, spatial coverage and the false-positive "
        "region fraction were defined as:",
    )
    spatial_equation = _new_paragraph(document, "")
    set_word_display_math(
        spatial_equation,
        r"C_{\mathrm{box}}=\frac{1}{N}\sum_{i=1}^{N}\max_j\operatorname{IoU}(G_i,P_j),"
        r"\quad f_{\mathrm{FP}}=1-\frac{\left|\left(\bigcup_jP_j\right)\cap"
        r"\left(\bigcup_iG_i\right)\right|}{\left|\bigcup_jP_j\right|}",
    )
    utility = _new_paragraph(
        document,
        "For within-modality descriptive ranking, the implemented utility was:",
    )
    utility_equation = _new_paragraph(document, "")
    set_word_display_math(
        utility_equation,
        r"U=0.35R_{\mathrm{PHI}}+0.25A_{\mathrm{field}}+0.15S_{\mathrm{text}}+"
        r"0.05C_{\mathrm{box}}+0.05S_{\mathrm{over}}+0.10S_{\mathrm{speed}}+"
        r"0.05M_{\mathrm{stab}}",
    )
    utility_definition = _new_paragraph(
        document,
        "Stext is one minus the mean normalized character and word error rates; Sover combines "
        "one minus the false-positive region fraction with one minus non-PHI area removed when "
        "canvas dimensions are available; Sspeed is the smaller of one and target latency "
        "divided by observed latency (0.50 s per video frame; 5.00 s per report); and Mstab "
        "discounts successful execution for positive resident-memory growth up to 256 MB. "
        "Utility is forced to zero when PHI-field recall is below 0.95 and is not a substitute "
        "for the unaggregated safety endpoints.",
    )
    _insert_after(
        outcomes,
        (spatial, spatial_equation, utility, utility_equation, utility_definition),
    )


def _reframe_release_control(document: DocumentObject) -> None:
    release = _starts(document, "Automated processing did not authorize release")
    set_word_math(
        release,
        (
            "Automated processing did not authorize release. Approval remained disabled until "
            "a reloadable anonymized artifact existed, integrity checks had passed, and a "
            "qualified reviewer had inspected the complete output. For structured tables, the "
            "immutable de-identified real input is ",
            Math(r"R^{(0)}"),
            "; a candidate released view ",
            Math(r"R_{\mathrm{rel}}"),
            " is the multiset union of retained real rows ",
            Math(r"R_{\mathrm{real}}"),
            " and synthetic rows ",
            Math(r"R_{\mathrm{syn}}"),
            ". The permitted reference actions are AddSynth, ModifySynth, Generalize, "
            "SuppressValue, and SuppressTuple; modification of real clinical values is "
            "prohibited. Privacy, sensitive-attribute, utility, provenance, and review "
            "conditions form the feasibility predicate. A separate local heuristic score "
            "ranks permitted successors; no global optimum or completeness claim is made. "
            "Synthetic provenance remains available to authorized analysts, and synthetic "
            "rows must not enter real-patient denominators.",
        ),
    )
    proposition = _new_paragraph(
        document,
        "Proposition (released-view multiplicity). For quasi-identifier tuple q, define:",
    )
    equation = _new_paragraph(document, "")
    set_word_display_math(
        equation,
        r"k_{\mathrm{rel}}(q)=\left|E_{QI}^{R_{\mathrm{rel}}}(q)\right|,\quad "
        r"k_{\mathrm{real}}(q)=\left|E_{QI}^{R_{\mathrm{rel}}}(q)\cap R_{\mathrm{real}}\right|",
    )
    implication = _new_paragraph(document, "")
    set_word_math(
        implication,
        (
            "The checker may establish ",
            Math(r"k_{\mathrm{rel}}(q)\ge k"),
            ", but ",
            Math(r"k_{\mathrm{rel}}(q)\ge k\;\nRightarrow\;k_{\mathrm{real}}(q)\ge k"),
            ". Synthetic insertion therefore repairs only the declared released-view "
            "predicate and provides no lower bound on the number of real subjects represented "
            "by that tuple. It is an experimental governance mechanism, not demonstrated "
            "re-identification protection.",
        ),
    )
    _insert_after(release, (proposition, equation, implication))


def _add_performance_caption(document: DocumentObject) -> None:
    heading = _exact(document, "4.1 Video-Level Anonymization Throughput and System Latency")
    heading.text = "4.1 Instrumentation verification"
    table = next(table for table in document.tables if table.cell(0, 0).text == "Performance metric")
    caption = _caption(
        document,
        "Table 2. Instrumented pipeline measures. These fields define the planned measurements; "
        "no comparative performance estimate is reported.",
    )
    table._tbl.addprevious(caption._p)
    _format_table(table, (2.15, 1.65, 2.70), numeric_from=3)


def _replace_integration_results(document: DocumentObject) -> None:
    heading = _exact(document, "4.2 OCR and De-Identification Quality Evaluation")
    heading.text = "4.2 Technical integration test and detector-development evaluation"
    overview = _starts(document, "The validation interface aggregates")
    overview.text = (
        "The four-frame endoscopy experiment tested execution of the corrected production "
        "path. It measured structured PHI-field extraction and spatial mask behavior, not "
        "residual readable PHI or independent-sample reliability. All frames came from one "
        "patient and one video; the unit is therefore one correlated integration fixture."
    )
    first_metric = _starts(document, "Residual OCR Match Count")
    last_context = _starts(document, "Human-in-the-loop annotations provide")
    paragraphs = document.paragraphs
    start = _index(paragraphs, first_metric)
    stop = _index(paragraphs, last_context)
    _remove_paragraphs(paragraphs[start : stop + 1])

    integration = _starts(document, "Post-fix production-path integration verification")
    integration.text = (
        "The V5 FrameOCR cascade and V6 FrameOCR plus the ONNX detector were applied to four "
        "frames containing first name, last name, date of birth, and examination date. Vision-"
        "language OCR was disabled. Results are descriptive fixture values; no confidence "
        "interval or patient-level performance estimate is available."
    )

    first_table_text = _exact(document, "Pipeline")
    caption_text = _starts(document, "Post-fix OCR integration results")
    paragraphs = document.paragraphs
    start = _index(paragraphs, first_table_text)
    stop = _index(paragraphs, caption_text)
    _remove_paragraphs(paragraphs[start : stop + 1])

    caption = _caption(
        document,
        "Table 3. Four-frame production-path integration test. All frames came from one patient "
        "and one video; values are descriptive, not independent validation estimates.",
    )
    table = _table(
        document,
        ("Pipeline", "Frames", "Field recall", "Field accuracy", "Box coverage", "FP-region fraction"),
        (
            ("V5 FrameOCR", "4", "1.000", "1.000", "0.748", "0.429"),
            ("V6 FrameOCR + ONNX", "4", "1.000", "1.000", "0.751", "0.800"),
        ),
        (2.05, 0.65, 0.95, 1.05, 0.95, 1.10),
    )
    operational_caption = _caption(
        document,
        "Table 4. Operational measurements from the same correlated four-frame fixture. "
        "Utility is defined in Section 3.6.",
    )
    operational = _table(
        document,
        ("Pipeline", "Mean latency (ms/frame)", "Mean utility score"),
        (
            ("V5 FrameOCR", "1,623.7", "0.851"),
            ("V6 FrameOCR + ONNX", "1,800.7", "0.835"),
        ),
        (3.25, 1.65, 1.60),
    )
    _insert_after(integration, (caption, table, operational_caption, operational))

    interpretation = _starts(document, "All eight sample–pipeline executions")
    interpretation.text = (
        "All eight frame–pipeline executions recovered the expected structured fields. This "
        "shows that the corrected code path executed on the fixture; it does not show that all "
        "readable PHI was removed. Adding the detector changed mean box coverage from 0.748 to "
        "0.751 and increased the false-positive region fraction from 0.429 to 0.800. The "
        "observed trade-off supports additive masking only under mandatory review."
    )

    caveat = _starts(document, "These values are integration evidence only")
    caveat.text = (
        "The scanned-report path was outside this passing result. A publishable clinical "
        "evaluation requires a frozen patient- or video-level test set spanning devices, "
        "rooms, PHI categories, negative frames, blur, occlusion, and low contrast; clustered "
        "confidence intervals; fixed-mask and OCR/detector baselines; temporal-ablation "
        "results; residual readable-PHI assessment; viewport-loss measurement; and manual-"
        "review failure rates."
    )

    detector_intro = _starts(document, "The refined pinned ONNX detector was executed")
    first_detector = _exact(document, "Modality")
    detector_summary = _starts(document, "Mean ground-truth box coverage was")
    paragraphs = document.paragraphs
    start = _index(paragraphs, first_detector)
    stop = _index(paragraphs, detector_summary) - 1
    _remove_paragraphs(paragraphs[start : stop + 1])
    detector_caption = _caption(
        document,
        "Table 5. Detector-development localization results at confidence 0.10 and matching "
        "IoU 0.50. Rows summarize annotated PHI instances, not independent patients.",
    )
    detector_table = _table(
        document,
        ("Modality", "Annotations", "Predictions", "Precision", "Recall", "F1", "Mean best IoU"),
        (
            ("CR", "14", "15", "0.933", "1.000", "0.966", "0.915"),
            ("DX", "18", "24", "0.750", "1.000", "0.857", "0.897"),
            ("MG", "3", "5", "0.600", "1.000", "0.750", "0.820"),
            ("Overall", "35", "44", "0.795", "1.000", "0.886", "0.897"),
        ),
        (0.85, 0.85, 0.85, 0.85, 0.75, 0.70, 1.15),
    )
    _insert_after(detector_intro, (detector_caption, detector_table))

    for malformed in list(document.tables):
        values = [cell.text.strip() for row in malformed.rows for cell in row.cells]
        if not values or values[0] in {
            "Performance metric",
            "Pipeline",
            "Modality",
            "Category",
            "Official action",
            "Element",
        }:
            continue
        parent = malformed._tbl.getparent()
        if parent is not None:
            parent.remove(malformed._tbl)


def _format_official_results(document: DocumentObject) -> None:
    intro = _starts(document, "The organizer validator indexed")
    intro.text = (
        "The official validator indexed all 23,921 exported DICOM files, reported zero "
        "missing-file batches, and completed 6,173,204 action checks. These checks are nested "
        "within files and rules and are not independent observations. The raw conformance rate "
        "was 75.56%. The organizer-defined 70% HIPAA, 20% DICOM Standard, and 10% Best Practice "
        "weights yielded 90.13%; this weighted score is neither a clinical accuracy estimate "
        "nor evidence that every identifier was removed."
    )
    category = next(table for table in document.tables if table.cell(0, 0).text == "Category")
    action = next(table for table in document.tables if table.cell(0, 0).text == "Official action")
    category_caption = _caption(
        document,
        "Table 6. Official MIDI-B weighted-category results. Weights were specified by the "
        "challenge organizers; the pooled pass rate treats correlated checks descriptively.",
    )
    action_caption = _caption(
        document,
        "Table 7. Official action-level DICOM results. Text-removal failures are potentially "
        "privacy-relevant; text-retention failures primarily indicate conformance or utility loss.",
    )
    category._tbl.addprevious(category_caption._p)
    action._tbl.addprevious(action_caption._p)
    _format_table(category, (1.35, 1.05, 1.05, 1.10, 0.85, 1.10))
    _format_table(action, (1.70, 1.15, 1.15, 1.20, 1.30))

    interpretation = _starts(document, "Privacy-critical checks were strongest")
    interpretation.text = (
        "Pixel hiding (35/35), date shifting (139,774/139,774), and patient-ID consistency "
        "(23,921/23,921) passed all applicable checks. However, the 125,101 failed text-removal "
        "checks are also privacy-relevant because retained descriptive attributes may contain "
        "identifying information; they require tag-level adjudication before any release claim. "
        "The 1,381,256 text-retention failures predominantly reflect conservative removal and "
        "therefore concern conformance and utility, while two pixels-retained failures quantify "
        "a small observed false-positive masking burden. The weighted score must be interpreted "
        "alongside these action-level failures."
    )


def _reframe_tabular_results(document: DocumentObject) -> None:
    heading = _exact(document, "4.4 Verification of Privacy-Preserving Data Publishing Frameworks")
    heading.text = "4.4 Structural check of the experimental tabular release control"
    opening = _starts(document, "The secondary-use tabular engine")
    opening.text = (
        "Deterministic dry runs exercised the declared state transitions and checker within the "
        "SensitiveMeta schema. They verify implementation structure only. No real-versus-"
        "synthetic distinguishability attack, auxiliary-information attack, disclosure-risk "
        "comparison, or utility calibration was performed; reduced re-identification risk is "
        "therefore not an empirical result of this paper."
    )
    score = _starts(document, "The reference optimization model balances")
    score.text = (
        "A local heuristic selection score ranks permitted successors after the feasibility "
        "predicate is checked. The score does not define a demonstrated global optimization "
        "procedure, and termination or predicate acceptance does not establish substantive "
        "privacy adequacy."
    )
    equation = document.paragraphs[_index(document.paragraphs, score) + 1]
    set_word_display_math(
        equation,
        r"L_{\mathrm{total}}=\alpha L_{\mathrm{size}}+\beta L_{\mathrm{sens}}+"
        r"\gamma L_{\mathrm{dist}}",
    )
    definitions = _starts(document, "where ")
    set_word_math(
        definitions,
        (
            "Here, ",
            Math(r"L_{\mathrm{size}}"),
            " measures synthetic insertion, ",
            Math(r"L_{\mathrm{sens}}"),
            " measures changes to synthetic sensitive values, and ",
            Math(r"L_{\mathrm{dist}}"),
            " is a prespecified utility discrepancy. Because feature weights, missing-value "
            "handling, joint distributions, distance scaling, and the acceptance threshold "
            "were not calibrated on a clinical release dataset, no utility or privacy-"
            "performance claim is made from this score.",
        ),
    )
    audit = _starts(document, "Audit decision: if a full-")
    audit.text = (
        "The checker accepts only a candidate satisfying every configured predicate. This is "
        "conditional soundness relative to the programmed checker, assuming correct "
        "canonicalization, finite successor generation, accurate real/synthetic labels, "
        "conservative numerical comparison, and integrity of protected reference data."
    )
    preliminary = _starts(document, "Preliminary dry runs are consistent")
    preliminary.text = (
        "Before this mechanism can support a release study, it requires comparison with "
        "suppression, generalization, aggregation, and non-release; a classifier attack on "
        "synthetic provenance; recomputation of real-subject class counts after removing "
        "distinguishable synthetic rows; and evaluation under plausible auxiliary knowledge."
    )


def _rewrite_limitations_and_conclusion(document: DocumentObject) -> None:
    legal = _starts(document, "The legal status of processed media")
    legal.text = legal.text.replace("access to control", "access-control")

    conclusion = _starts(document, "Endoscopic video de-identification is best understood")
    conclusion.text = (
        "This work defines a reproducible, fail-closed workflow for local PHI localization, "
        "metadata transformation, pseudonymous linkage, DICOM export, and human-controlled "
        "release. Temporal metadata accumulation and device-aware masking are implementation "
        "features intended to preserve clinical content while reducing identifier exposure. "
        "Pseudonyms remain personal data where re-identification is reasonably possible, and "
        "synthetic padding changes only a released-view predicate rather than the number of "
        "real people represented."
    )
    evidence = _starts(document, "Absolute anonymity is not inferred")
    evidence.text = (
        "The measured evidence is narrower than a clinical validation claim. The four-frame "
        "endoscopy test demonstrates correct execution on one correlated fixture. The "
        "35-annotation detector analysis is a development result on a refined validation split. "
        "The exhaustive 23,921-file DICOM run demonstrates export and conformance behavior, but "
        "its 90.13% weighted score coexists with 125,101 failed text-removal checks and does not "
        "establish universal DICOM safety. Submission as a clinical performance study requires "
        "an independent patient- and video-level evaluation across sites, devices, PHI types, "
        "and adverse image conditions, with residual readable-PHI and complete-redaction rates "
        "as primary safety endpoints."
    )


def _clean_references(document: DocumentObject) -> None:
    unused_starts = (
        "Denning, D., & Lunt, T.",
        "FIDO Alliance.",
        "HashiCorp.",
        "International Organization for Standardization.",
        "National Institute of Standards and Technology. (2007).",
        "ZER-0-NE.",
    )
    remove: list[Paragraph] = []
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(unused_starts):
            remove.append(paragraph)
            continue
        if " Cited by:" in paragraph.text:
            paragraph.text = paragraph.text.split(" Cited by:", 1)[0]
        if "healthcare13202594X" in paragraph.text:
            paragraph.text = paragraph.text.replace("healthcare13202594X", "healthcare13202594")
        if "https://doi.org/10.20965/jaciii.2024.p0000" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                " https://doi.org/10.20965/jaciii.2024.p0000", ""
            )
    _remove_paragraphs(remove)


def revise(source: Path, output: Path) -> None:
    if source.resolve() == output.resolve():
        raise ValueError("Source and output paths must differ")
    document = Document(str(source))
    _rewrite_abstract_and_focus(document)
    _add_threat_model(document)
    _clarify_methods_and_endpoints(document)
    _reframe_release_control(document)
    _add_performance_caption(document)
    _replace_integration_results(document)
    _format_official_results(document)
    _reframe_tabular_results(document)
    _rewrite_limitations_and_conclusion(document)
    _clean_references(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Apply the manuscript's main reviewer corrections")
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    return parser


def main() -> int:
    arguments = _parser().parse_args()
    revise(arguments.source, arguments.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
