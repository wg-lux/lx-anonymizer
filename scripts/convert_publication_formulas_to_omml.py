from __future__ import annotations

import argparse
import copy
import subprocess
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Iterable
from zipfile import ZipFile

from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
import lxml.etree as etree


MATH_NAMESPACE = "http://schemas.openxmlformats.org/officeDocument/2006/math"


@dataclass(frozen=True)
class Math:
    latex: str


Part = str | Math


def set_word_math(paragraph: Paragraph, parts: Iterable[Part]) -> None:
    """Replace a paragraph with mixed text and native inline Word equations."""
    _set_mixed(paragraph, parts)


def set_word_display_math(paragraph: Paragraph, latex: str) -> None:
    """Replace a paragraph with one centered native Word equation."""
    _set_display(paragraph, latex)


def _pandoc_math(latex: str, *, display: bool = False) -> etree._Element:
    delimiter = "$$" if display else "$"
    markdown = f"{delimiter}{latex}{delimiter}\n"
    completed = subprocess.run(
        ["pandoc", "-f", "markdown", "-t", "docx", "-o", "-"],
        input=markdown.encode(),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=False,
    )
    if completed.returncode != 0:
        raise RuntimeError(completed.stderr.decode(errors="replace"))
    with ZipFile(BytesIO(completed.stdout)) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    local_name = "oMathPara" if display else "oMath"
    matches = root.xpath(f".//m:{local_name}", namespaces={"m": MATH_NAMESPACE})
    if len(matches) != 1:
        raise ValueError(f"Expected one {local_name} element for {latex!r}, found {len(matches)}")
    return copy.deepcopy(matches[0])


def _set_mixed(paragraph: Paragraph, parts: Iterable[Part]) -> None:
    paragraph.clear()
    paragraph.style = "Normal"
    for part in parts:
        if isinstance(part, Math):
            paragraph._p.append(_pandoc_math(part.latex))
        else:
            paragraph.add_run(part)


def _set_display(paragraph: Paragraph, latex: str) -> None:
    paragraph.clear()
    paragraph.style = "Normal"
    paragraph._p.append(_pandoc_math(latex, display=True))


def _remove(paragraphs: Iterable[Paragraph]) -> None:
    for paragraph in paragraphs:
        element = paragraph._p
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def _exact(document: DocumentObject, text: str) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == text]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def _index(paragraphs: list[Paragraph], target: Paragraph) -> int:
    for index, paragraph in enumerate(paragraphs):
        if paragraph._p is target._p:
            return index
    raise ValueError(f"Paragraph not found: {target.text!r}")


def _replace_related_work_math(document: DocumentObject) -> None:
    first = _exact(
        document,
        "-Anonymity: Dictates that each sequence of quasi-identifiers must be "
        "indistinguishable from at least",
    )
    last = _exact(document, "diversity have proven weaknesses. (Gadotti, 2024)")
    paragraphs = document.paragraphs
    start = _index(paragraphs, first)
    stop = _index(paragraphs, last)
    old = paragraphs[start : stop + 1]

    replacements: tuple[tuple[Part, ...], ...] = (
        (
            "k-anonymity requires each released quasi-identifier tuple to occur in at "
            "least ",
            Math("k"),
            " records, so that a record shares its tuple with at least ",
            Math("k-1"),
            " other released records [Sweeney, 2002].",
        ),
        (
            "l-diversity addresses homogeneity attacks by requiring at least ",
            Math("l"),
            " well-represented sensitive values in each equivalence class "
            "[Machanavajjhala et al., 2007].",
        ),
        (
            "Total-variation t-closeness bounds the discrepancy between the sensitive-"
            "attribute distribution in a class and a declared global reference by ",
            Math("t"),
            " [Li et al., 2007].",
        ),
        (
            "Differential privacy, commonly parameterized by ",
            Math(r"\varepsilon"),
            ", provides query-level guarantees but does not directly preserve record-level "
            "linkage between videos, annotations, and reports [Dwork and Roth, 2014]. These "
            "models address different risks, and neither k-anonymity nor l-diversity is "
            "sufficient against all linkage or attribute-disclosure attacks [Gadotti, 2024].",
        ),
    )
    for paragraph, parts in zip(old, replacements, strict=False):
        _set_mixed(paragraph, parts)
    _remove(old[len(replacements) :])


def _replace_method_metrics(document: DocumentObject) -> None:
    paragraph = next(
        item
        for item in document.paragraphs
        if item.text.startswith("For box localization, a prediction was a true positive")
    )
    _set_mixed(
        paragraph,
        (
            "For box localization, a prediction was a true positive when it matched an "
            "annotated PHI box at ",
            Math(r"\operatorname{IoU} \ge 0.50"),
            ". We report precision ",
            Math(r"P=\frac{TP}{TP+FP}"),
            ", recall ",
            Math(r"R=\frac{TP}{TP+FN}"),
            ", and ",
            Math(r"F_1=\frac{2PR}{P+R}"),
            ", together with mean best IoU and mean annotated-box coverage. Results were "
            "stratified by modality (CR, DX, and MG). For DICOM validation, action pass "
            "rate was ",
            Math(r"r=\frac{\mathrm{pass}}{\mathrm{pass}+\mathrm{fail}}"),
            ". The raw pass rate pooled all actions. The official challenge score was ",
            Math(
                r"S=0.70r_{\mathrm{HIPAA}}+0.20r_{\mathrm{DICOM}}+"
                r"0.10r_{\mathrm{BestPractice}}"
            ),
            ". Counts are reported without sampling confidence intervals because the "
            "validator exhaustively tested the exported collection; this does not remove "
            "uncertainty about generalization to other institutions, devices, or populations.",
        ),
    )

    release = next(
        item
        for item in document.paragraphs
        if item.text.startswith("Automated processing did not authorize release.")
    )
    text = release.text
    before_first, remainder = text.split(" k,", 1)
    before_second, after_second = remainder.split(" met k,", 1)
    _set_mixed(
        release,
        (
            before_first,
            " ",
            Math("k"),
            ",",
            before_second,
            " met ",
            Math("k"),
            ",",
            after_second,
        ),
    )


def _replace_throughput_math(document: DocumentObject) -> None:
    first = _exact(
        document,
        "The evaluation suite tracks resource allocation metrics at execution time, "
        "capturing total lifecycle time",
    )
    tail = _exact(
        document,
        ". Frame throughput rates are dynamically derived by pulling frame counts directly "
        "from the media database schema:",
    )
    paragraphs = document.paragraphs
    start = _index(paragraphs, first)
    stop = _index(paragraphs, tail)
    old = paragraphs[start : stop + 1]
    _set_mixed(
        first,
        (
            "The evaluation suite records total lifecycle time ",
            Math(r"T_{\mathrm{total}}"),
            ", staging time ",
            Math(r"T_{\mathrm{stage}}"),
            ", and anonymization time ",
            Math(r"T_{\mathrm{anon}}"),
            ". Frame throughput is derived from decoded frame count and elapsed "
            "anonymization time.",
        ),
    )
    _remove(old[1:])

    table_start = _exact(document, "Performance Metric Captured")
    table_stop = _exact(document, "Defines capability for near-line processing vs.\u00a0batch offloading")
    paragraphs = document.paragraphs
    start = _index(paragraphs, table_start)
    stop = _index(paragraphs, table_stop)
    _remove(paragraphs[start : stop + 1])

    for legacy in list(document.tables):
        values = [cell.text.strip() for row in legacy.rows for cell in row.cells]
        nonempty = [value for value in values if value]
        if nonempty == ["T_{anon}"]:
            parent = legacy._tbl.getparent()
            if parent is not None:
                parent.remove(legacy._tbl)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ("Performance metric", "Measurement", "Operational interpretation")
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.bold = True

    rows: tuple[tuple[Part, str, str], ...] = (
        (
            "Ingestion staging latency",
            "staging_seconds",
            "Filesystem allocation, metadata parsing, and container verification",
        ),
        (
            Math(r"T_{\mathrm{anon}}"),
            "anonymizer_seconds",
            "Processor-aware masking and configured local OCR fallback",
        ),
        (
            "Memory footprint delta",
            "RSS delta; CPU time",
            "Resident-memory change and processing cost",
        ),
        (
            "Throughput",
            "frames per second",
            "Decoded frames divided by elapsed anonymization time",
        ),
    )
    for metric, measurement, interpretation in rows:
        cells = table.add_row().cells
        _set_mixed(cells[0].paragraphs[0], (metric,))
        cells[1].text = measurement
        cells[2].text = interpretation
    first._p.addnext(table._tbl)


def _replace_release_score_math(document: DocumentObject) -> None:
    k_paragraph = next(
        item
        for item in document.paragraphs
        if item.text.startswith("The reference optimization model balances")
    )
    prefix, suffix = k_paragraph.text.split("k (", 1)
    _set_mixed(k_paragraph, (prefix, Math("k"), " (", suffix))

    score = _exact(
        document,
        "L_{Total} = \\alpha\\, L_{size} + \\beta L_{sens} + \\gamma L_{dist},",
    )
    _set_display(
        score,
        r"L_{\mathrm{total}}=\alpha L_{\mathrm{size}}+\beta L_{\mathrm{sens}}+"
        r"\gamma L_{\mathrm{dist}}",
    )

    where = _exact(document, "where")
    size = _exact(document, "L_{size}")
    size_description = _exact(document, "measures synthetic padding volume,")
    sensitivity = _exact(document, "L_{sens}")
    sensitivity_description = _exact(
        document, "measures modifications to synthetic sensitive values, and"
    )
    distance = _exact(document, "L_{dist}")
    distance_description = next(
        item
        for item in document.paragraphs
        if item.text.startswith("measures the pre-specified distributional discrepancy.")
    )
    _set_mixed(
        where,
        (
            "where ",
            Math(r"L_{\mathrm{size}}"),
            " measures synthetic padding volume, ",
            Math(r"L_{\mathrm{sens}}"),
            " measures changes to synthetic sensitive values, and ",
            Math(r"L_{\mathrm{dist}}"),
            " measures the prespecified distributional discrepancy. In the conservative "
            "reference algorithm, the successor relation prohibits modification of real "
            "clinical rows rather than approximating that prohibition with an infinite "
            "penalty.",
        ),
    )
    _remove(
        (
            size,
            size_description,
            sensitivity,
            sensitivity_description,
            distance,
            distance_description,
        )
    )

    audit = next(
        item for item in document.paragraphs if item.text.startswith("Audit decision: if a full-")
    )
    paragraphs = document.paragraphs
    audit_index = _index(paragraphs, audit)
    audit_tail = next(
        item
        for item in paragraphs[audit_index:]
        if item.text.startswith(", the algorithm calculates its deficit")
    )
    audit_stop = _index(paragraphs, audit_tail)
    _set_mixed(
        audit,
        (
            "Audit decision: if a full-",
            Math("QI"),
            " class has count ",
            Math("c(QI,q)<k"),
            ", the algorithm calculates its deficit and proposes a permitted repair. If no "
            "deficit and no configured distributional or utility violation remains, the "
            "candidate satisfies the release predicate. The output is described as "
            "predicate-compliant, not categorically safe or anonymous.",
        ),
    )
    _remove(paragraphs[audit_index + 1 : audit_stop + 1])

    preliminary = next(
        item
        for item in document.paragraphs
        if item.text.startswith("Preliminary dry runs are consistent")
    )
    paragraphs = document.paragraphs
    preliminary_index = _index(paragraphs, preliminary)
    variability = next(
        item
        for item in paragraphs[preliminary_index:]
        if item.text.endswith("thresholds, and run-to-run variability.")
    )
    variability_index = _index(paragraphs, variability)
    _set_mixed(
        preliminary,
        (
            "Preliminary dry runs are consistent with the expected qualitative trade-off "
            "that larger ",
            Math("k"),
            " can require more synthetic padding in sparse full-",
            Math("QI"),
            " classes. The magnitude of padding and distributional distortion cannot be "
            "claimed until the evaluated tables, class counts, synthetic-row proportions, ",
            Math(r"D_{\mathrm{util}}"),
            " values, total-variation distances, thresholds, and run-to-run variability "
            "are reported.",
        ),
    )
    _remove(paragraphs[preliminary_index + 1 : variability_index + 1])


def convert(source: Path, output: Path) -> None:
    if source.resolve() == output.resolve():
        raise ValueError("Source and output paths must differ")
    document = Document(str(source))
    _replace_related_work_math(document)
    _replace_method_metrics(document)
    _replace_throughput_math(document)
    _replace_release_score_math(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Convert publication formulas to Word OMML")
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    return parser


def main() -> int:
    arguments = _parser().parse_args()
    convert(arguments.source, arguments.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
